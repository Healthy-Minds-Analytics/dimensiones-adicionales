from io import BytesIO
import os
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import glob
from datetime import datetime
import json
import random
from copy import deepcopy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def seleccionar_csv(ruta):
    """Busca archivos CSV en la carpeta de la ruta proporcionada.
       Si solo hay uno, lo devuelve, y si hay más de uno
       escoge el más reciente."""
    
    patron_busqueda = os.path.join(ruta, "*.csv")
    archivos_csv = glob.glob(patron_busqueda)
    
    if not archivos_csv:
        print(f"No se encontró ningún archivo CSV en la carpeta {ruta}.")
        return None

    if len(archivos_csv) == 1:
        return archivos_csv[0]

    # Si hay varios, escogemos el más reciente (mayor fecha de modificación)
    archivo_mas_reciente = max(archivos_csv, key=os.path.getmtime)
    return archivo_mas_reciente

def replace_bookmark_pair(doc, pair):
    bookmark_name, replacement = pair
    replacement = "" if replacement is None else str(replacement)
    found = False

    def replace_in_element(el):
        nonlocal found
        for child in el:
            # 1) Localizamos el bookmarkStart
            if child.tag.endswith('bookmarkStart') and child.get(qn('w:name')) == bookmark_name:
                found = True

                # 2) Buscamos el run (<w:r>) inmediato que contiene el <w:t>
                run_elem = child.getnext()
                while run_elem is not None and not run_elem.tag.endswith('r'):
                    run_elem = run_elem.getnext()
                if run_elem is None:
                    return

                text_elem = run_elem.find(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
                )
                if text_elem is None:
                    return

                # 3) Si no hay saltos de línea, uso la lógica antigua: sólo reemplazo el marcador
                if '\n' not in replacement:
                    # Reemplazo el contenido del <w:t>
                    text_elem.text = str(replacement)
                    # Elimino todos los nodos hasta encontrar el bookmarkEnd
                    sib = run_elem.getnext()
                    while sib is not None and not sib.tag.endswith('bookmarkEnd'):
                        to_remove = sib
                        sib = sib.getnext()
                        to_remove.getparent().remove(to_remove)
                    return

                # 4) Si hay saltos de línea, uso la lógica de “múltiples párrafos”
                #    Primero, localizo el párrafo padre para clonar estilo
                p_elem = run_elem.getparent()
                # Borro todos los runs viejos
                for r in p_elem.findall(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'
                ):
                    p_elem.remove(r)

                # 5) Divido el replacement en líneas y genero párrafos
                lines = str(replacement).split('\n')

                # Primera línea en el párrafo original
                first_r = OxmlElement('w:r')
                first_t = OxmlElement('w:t')
                first_t.text = lines[0]
                first_r.append(first_t)
                p_elem.append(first_r)

                # Líneas extra → clono párrafo (preserva bullets/numPr) y añado nuevo texto
                prev_p = p_elem
                for line in lines[1:]:
                    new_p = deepcopy(p_elem)  # hereda <w:pPr> (bullet o estilo)
                    # limpio runs clonados
                    for r in new_p.findall(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'
                    ):
                        new_p.remove(r)
                    # inserto el run+t con la nueva línea
                    nr = OxmlElement('w:r')
                    nt = OxmlElement('w:t')
                    nt.text = line
                    nr.append(nt)
                    new_p.append(nr)
                    prev_p.addnext(new_p)
                    prev_p = new_p

                return

            # 6) Si no es este marcador, sigo bajando recursivamente
            replace_in_element(child)

    replace_in_element(doc._element)
    if not found:
        print(f"Marcador '{bookmark_name}' no encontrado")

def obtenerRespuestas(dataframe, inicio, fin):
    """
    Genera un diccionario con el conteo de cada respuesta por pregunta en un DataFrame, 
    construyendo las claves en el formato 'PREGUNTA_X_Y'.

    Parámetros
    ----------
    dataframe : pd.DataFrame
        DataFrame donde cada columna representa una pregunta y cada fila registra las 
        respuestas de un encuestado. Los valores del DataFrame deben coincidir con las 
        posibles respuestas definidas en mapa_respuestas.
    mapa_respuestas : dict
        Diccionario que mapea las respuestas posibles (claves) a valores numéricos (valores). 
        Ejemplo: {"Muy de acuerdo": 5, "De acuerdo": 4, ...}

    Proceso de la función
    ---------------------
    1. Se inicializa un diccionario vacío (conteo_respuestas) para acumular los conteos.
    2. Se extraen las listas de respuestas y valores del diccionario mapa_respuestas:
       - respuestas_posibles (lista de claves).
       - valores_posibles (lista de valores numéricos).
    3. Se recorren las columnas del DataFrame (enumerate(dataframe.columns, start=1)):
       - 'i' representará el índice o número de pregunta.
       - 'pregunta' será el nombre de la columna.
       - Para cada pregunta, se obtiene un conteo de respuestas usando `dataframe[pregunta].value_counts()`.
    4. Se rellena el diccionario de conteos con cada respuesta posible para esa pregunta:
       - Por cada par (respuesta, valor) en zip(respuestas_posibles, valores_posibles):
         - Se construye la clave: f"PREGUNTA_{i}_{valor}".
         - Se asigna el número de ocurrencias de dicha 'respuesta' en la serie `conteo`. 
           Si la respuesta no existe en la columna, se asigna 0.

    Valor de retorno
    ----------------
    dict
        Diccionario que asocia claves en el formato 'PREGUNTA_X_Y' 
        (X = número de pregunta, Y = valor numérico de la respuesta) 
        con el conteo de cuántas veces apareció esa respuesta en dicha pregunta.
        Ejemplo de clave: "PREGUNTA_3_5".

    Ejemplo de uso
    --------------
    >>> import pandas as pd
    >>> df = pd.DataFrame({
    ...     "Pregunta 1": ["Muy de acuerdo", "De acuerdo", "Muy de acuerdo"],
    ...     "Pregunta 2": ["En desacuerdo", "Muy de acuerdo", "De acuerdo"]
    ... })
    >>> mapa = {"Muy de acuerdo": 5, "De acuerdo": 4, "Ni de acuerdo ni en desacuerdo": 3,
    ...         "En desacuerdo": 2, "Muy en desacuerdo": 1}
    >>> conteos = obtenerRespuestas(df, mapa)
    >>> print(conteos)
    {
       "PREGUNTA_1_5": 2,
       "PREGUNTA_1_4": 1,
       "PREGUNTA_1_3": 0,
       "PREGUNTA_1_2": 0,
       "PREGUNTA_1_1": 0,
       "PREGUNTA_2_5": 1,
       "PREGUNTA_2_4": 1,
       "PREGUNTA_2_3": 0,
       "PREGUNTA_2_2": 1,
       "PREGUNTA_2_1": 0
    }
    """
    conteo_respuestas = {}
    valores_posibles = range(inicio, fin)
    
    for i, pregunta in enumerate(dataframe.columns, start=1):
        # Contar respuestas para la pregunta
        conteo = dataframe[pregunta].value_counts()
        
        for valor in valores_posibles:
            clave = f"PREGUNTA_{i}_{valor}"
            conteo_respuestas[clave] = conteo.get(valor, 0)  # Obtener el conteo o 0 si no aparece
    
    return conteo_respuestas

def calcularValores(respuestas_dim: pd.DataFrame) -> pd.DataFrame:
    """
    Devuelve un DataFrame con la media y la desviación estándar
    por cada dimensión (como filas), ya redondeadas a 2 decimales.
    
    Índices: nombre de la dimensión.
    Columnas: ['mean', 'std'].
    """
    stats = respuestas_dim.agg(['mean', 'std']).T
    return stats.round(2)

def df_a_reemplazos(df_stats: pd.DataFrame) -> dict:
    reemplazos = {}
    for dim, row in df_stats.iterrows():
        reemplazos[f"MEDIA_{dim}"] = row['mean']
        reemplazos[f"STD_{dim}"] = row['std']
    return reemplazos

def escogerMedidas(estadisticas: pd.DataFrame, ruta_medidas, limite=10) -> dict:
    estadisticas_corregidas = estadisticas.copy()
    estadisticas_corregidas.loc[['FISICAS', 'SOCIALES', 'PSICOLOGICAS'], 'mean'] *= 3

    # Buscar las dimensiones (filas) de estadisticas cuya media sea mayor que el límite indicado
    alertas = estadisticas_corregidas[estadisticas_corregidas['mean'] > limite]

    dims = list(alertas.index)

    if len(dims) < 2:
        dims = estadisticas_corregidas['mean'].nlargest(2).index.tolist()


    ficheros = {
        "CARACTERISTICAS_TAREA": 'caracteristicas_tarea',
        "ORGANIZACION":'organizacion',
        "TEDIO": 'tedio',
        "CANSANCIO_EMOCIONAL": 'cansancio_emocional',
        "DESPERSONALIZACION": 'despersonalizacion',
        "REALIZACION_PERSONAL": 'realizacion_personal',
        'FISICAS': 'consecuencias_fisicas',
        'SOCIALES': 'consecuencias_sociales',
        'PSICOLOGICAS': 'consecuencias_psicologicas'
        
    }

    parrafos = []
    for dim in dims:
        base = ficheros.get(dim, '')
        if not base:
            # no hay JSON asociado: lo ignoramos
            print(f"No hay fichero asociado para la dimensión {dim!r}, por lo que no se proponen medidas para dicha dimensión.")
            continue

        fichero = os.path.join(ruta_medidas, f'{base}.json')
        if not os.path.exists(fichero):
            print(f"El fichero {fichero} no existe")
            continue
        
        # Abrir archivo JSON
        with open(fichero, 'r', encoding='utf-8') as f:
            data = json.load(f)
        if not data:
            print(f"El fichero {fichero} de la dimensión {dim!r} está vacío")
            continue
        
        try:
            lista_medidas = next(iter(data.values()))
        except StopIteration:
            continue

        if not lista_medidas:
            continue

        #print(lista_medidas)
        # Escoge una medida al azar
        medida = random.choice(lista_medidas)

        parrafos.append(f"{medida}")

    texto_final = "\n".join(parrafos)

    return {'MEDIDAS': texto_final}


def generarWord(plantilla_doc, informe, carpeta_informes, reemplazos):
    """
    Crea un documento de Word a partir de una plantilla, reemplazando cada marcador 
    (clave) del diccionario `reemplazos` por su valor correspondiente.

    Parámetros
    ----------
    plantilla_doc : str
        Ruta al archivo .docx que sirve de plantilla.
    informe : dict
        Nombre del informe a generar.
    carpeta_informes : str
        Ruta a la carpeta donde se guardará el informe.
    reemplazos : dict
        Diccionario cuyas claves son nombres de marcador y cuyos valores son 
        los textos que se insertarán en dichos marcadores.

    Genera
    ------
    Informe_Satisfaccion_Generado.docx
        Un archivo de Word con todos los marcadores reemplazados.
    """
    # Carpeta para guardar los informes
    if not os.path.exists(carpeta_informes):
        os.makedirs(carpeta_informes)
    
    doc = Document(plantilla_doc)

    # Aplicar reemplazos usando map
    list(map(lambda pair: replace_bookmark_pair(doc, pair), reemplazos.items()))

    output_doc = os.path.join(carpeta_informes, f"Informe_{informe}_{reemplazos['NOMBRE_EMPRESA']}.docx")
    doc.save(output_doc)
    
    print(f"Informe generdo correctamente. Cierre esta ventana y vaya a {output_doc}")

def generarWord_bytes(plantilla_doc: str, reemplazos: dict) -> bytes:
    doc = Document(plantilla_doc)

    list(map(lambda pair: replace_bookmark_pair(doc, pair), reemplazos.items()))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def calcular_metricas_estadisticas(respuestas_agrupadas: pd.DataFrame) -> dict:
    """
    Calcula estadísticas para cada dimensión y devuelve un diccionario de diccionarios.
    
    Estructura del retorno:
    {
        "Satisfaccion_General": {
            "media": ...,
            "std": ...,
            ...
        },
        ...
    }
    """
    metricas = {}

    dimensiones = [
        "Satisfaccion_General",
        "Satisfaccion_Intrinseca",
        "Satisfaccion_Extrinseca"
    ]

    for dimension in dimensiones:
        serie = respuestas_agrupadas[dimension]

        metricas[dimension] = {
            "media": round(serie.mean(), 2),
            "std": round(serie.std(), 2),
            "min": int(serie.min()),
            "max": int(serie.max()),
            "p25": round(serie.quantile(0.25), 2),
            "p75": round(serie.quantile(0.75), 2),
            "count": int(serie.count())
        }

    return metricas

def generar_informe_burnout(csv_source, empresa, invitados, limite=10) -> bytes:
    ruta_script = os.path.dirname("./Burnout/")
    carpeta_plantillas = os.path.join(ruta_script, "Plantillas")
    carpeta_medidas = os.path.join(ruta_script, "Medidas")
    ruta_config = os.path.join(ruta_script, "Dimensiones_CBB.json")

    # Carga de la configuración de dimensiones del CBB
    ruta_config = os.path.join(ruta_script, 'Dimensiones_CBB.json')
    with open(ruta_config, 'r', encoding='utf-8') as f:
        config = json.load(f)

    # Leer CSV (puede ser filepath o UploadedFile)
    if hasattr(csv_source, "read"):
        respuestas = pd.read_csv(csv_source, sep=None, engine="python")
    else:
        respuestas = pd.read_csv(csv_source, sep=None, engine="python")
    preguntas = list(respuestas.columns)

    # Mapa de respuestas CBB
    # Se normaliza todo a minúsculas y sin espacios sobrantes
    mapa_respuestas = {
    # Escala “nada” → “mucho” (ítems 13,15,16,17,19)
    "nada": 1,
    "muy poco": 2,
    "algo": 3,
    "bastante": 4,
    "mucho": 5,
    # Escala “en ninguna ocasión” → “en la mayoría de ocasiones” (ítems 1,2,3,4,7,8,14,20,21)
    "en ninguna ocasión": 1,
    "raramente": 2,
    "algunas veces": 3,
    "frecuentemente": 4,
    "en la mayoría de ocasiones": 5,
    # Escala “totalmente en desacuerdo” → “totalmente de acuerdo” (ítems 5,6,10,11,12)
    "totalmente en desacuerdo": 1,
    "en desacuerdo": 2,
    "indeciso": 3,
    "de acuerdo": 4,
    "totalmente de acuerdo": 5,
    # Escala “nunca” → “siempre” (ítems 9,18)
    "nunca": 1,
    "siempre": 5
    }

    # Conversión de texto a números
    '''respuestas_convertidas = respuestas.applymap(lambda x: mapa_respuestas.get(x.strip().lower(), x) if isinstance(x, str) else x)
    respuestas_convertidas = respuestas.applymap(lambda x: mapa_respuestas.get(x, x))'''

    respuestas_convertidas = respuestas.map(lambda x: mapa_respuestas.get(x.strip().lower(), x.strip().lower()) if isinstance(x, str) else x)

    # Agrupamiento por dimensión
    respuestas_agrupadas = pd.DataFrame(index=respuestas.index)

    for bloque_nombre, bloque in config.items():
        for dim_nombre, info in bloque.items():
            items = info['items']
            # Convertimos la lista de índices 1-based en nombres de columna
            cols = [preguntas[i-1] for i in items]
            respuestas_agrupadas[dim_nombre] = respuestas_convertidas[cols].sum(axis=1)


    # Cálculo de la participación
    informacion = {
        "NOMBRE_EMPRESA": empresa,
        "PARTICIPACION": round(len(respuestas)/invitados*100, 2) if invitados>0 else 0
    }

    estadisticas = calcularValores(respuestas_agrupadas)
    calculos = df_a_reemplazos(estadisticas)
    conteo_respuestas = obtenerRespuestas(respuestas, 1, 6)
    medidas = escogerMedidas(estadisticas, carpeta_medidas, limite)

    reemplazos = informacion | calculos | conteo_respuestas | medidas
    plantilla = os.path.join(carpeta_plantillas, "plantilla_burnout.docx")

    doc = Document(plantilla)
    list(map(lambda pair: replace_bookmark_pair(doc, pair), reemplazos.items()))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
