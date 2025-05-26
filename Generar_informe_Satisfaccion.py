from io import BytesIO
import os
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn 
import glob
from datetime import datetime
import json
import random

def seleccionar_csv(ruta):
    """Busca archivos CSV en la carpeta de la ruta proporcionada.
       Si solo hay uno, lo devuelve, y si hay más de uno
       escoge el más reciente."""
    
    patron_busqueda = os.path.join(ruta, "*.csv")
    archivos_csv = glob.glob(patron_busqueda)
    
    if not archivos_csv:
        print(f"No se encontró ningún archivo CSV en la carpeta {ruta}.")
        return None

    if len(archivos_csv) == 1:
        return archivos_csv[0]

    # Si hay varios, escogemos el más reciente (mayor fecha de modificación)
    archivo_mas_reciente = max(archivos_csv, key=os.path.getmtime)
    return archivo_mas_reciente


def replace_bookmark_pair(doc, pair):
    """
    Reemplaza el contenido asociado a un marcador específico en un documento Word (python-docx),
    recorriendo todo su árbol XML, incluidos cuadros de texto y demás estructuras anidadas.

    Parámetros
    ----------
    doc : docx.Document
        Objeto Document proporcionado por la librería python-docx. Representa el documento 
        donde se realizará la búsqueda y reemplazo.
    pair : tuple
        Tupla que contiene (bookmark_name, replacement).
        
        - bookmark_name (str): Nombre del marcador a localizar en el documento.
        - replacement (str): Texto o valor que se asignará en sustitución del contenido 
          hallado dentro del marcador.

    Comportamiento
    -------------
    1. Se define una función interna `replace_in_element(element)` que:
       - Recorre recursivamente cada uno de los subelementos del XML del documento.
       - Si encuentra un 'bookmarkStart' cuyo atributo 'w:name' coincida con bookmark_name:
         - Marca la variable `found` como True.
         - Avanza sobre los elementos hermanos (next_sibling) del marcador hasta localizar un 
           run (`<w:r>`) que contenga un elemento texto (`<w:t>`).
         - Reemplaza el contenido de `<w:t>` con la cadena `replacement`.
         - Luego elimina (en caso de existir) todos los elementos hermanos siguientes 
           hasta toparse con un 'bookmarkEnd' (indica el fin del marcador).
         - Termina el proceso tras el primer reemplazo exitoso.
       - Continúa explorando recursivamente el resto de elementos si no se ha encontrado el marcador.

    2. La función principal `replace_bookmark_pair(doc, pair)`:
       - Toma la raíz (`doc._element`) y la recorre llamando a `replace_in_element`.
       - Si, al finalizar el recorrido, la variable `found` sigue en False, 
         imprime un aviso por consola indicando que el marcador no se encontró.

    Notas
    ----
    - Este método modifica el documento en memoria: al finalizar, conviene llamar a `doc.save(...)`
      para persistir los cambios en un archivo.
    - La función solo realiza un reemplazo por marcador. Si un marcador aparece varias veces 
      con el mismo nombre, solo se reemplazará la primera aparición que se halle al recorrer el XML.
    - El proceso recursivo permite hallar el marcador aunque esté dentro de cuadros de texto, 
      tablas u otras secciones anidadas del documento.

    Ejemplo de uso
    --------------
    >>> from docx import Document
    >>> doc = Document("mi_documento.docx")
    >>> replace_bookmark_pair(doc, ("MI_MARKER", "Nuevo contenido"))
    >>> doc.save("mi_documento_modificado.docx")
    """
    bookmark_name, replacement = pair
    found = False

    def replace_in_element(element):
        nonlocal found
        for child in element:
            if child.tag.endswith('bookmarkStart') and child.get(qn('w:name')) == bookmark_name:
                found = True
                next_sibling = child.getnext()
                while next_sibling is not None:
                    if next_sibling.tag.endswith('r'):
                        text_element = next_sibling.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                        if text_element is not None:
                            text_element.text = str(replacement)
                            following_sibling = next_sibling.getnext()
                            while following_sibling is not None and not following_sibling.tag.endswith('bookmarkEnd'):
                                parent = following_sibling.getparent()
                                parent.remove(following_sibling)
                                following_sibling = next_sibling.getnext()
                            return
                    next_sibling = next_sibling.getnext()
            replace_in_element(child)  # Llamada recursiva

    replace_in_element(doc._element)

    if not found:
        print(f"Marcador '{bookmark_name}' no encontrado")


def obtenerRespuestas(dataframe, inicio, fin):
    """
    Genera un diccionario con el conteo de cada respuesta por pregunta en un DataFrame, 
    construyendo las claves en el formato 'PREGUNTA_X_Y'.

    Parámetros
    ----------
    dataframe : pd.DataFrame
        DataFrame donde cada columna representa una pregunta y cada fila registra las 
        respuestas de un encuestado. Los valores del DataFrame deben coincidir con las 
        posibles respuestas definidas en mapa_respuestas.
    mapa_respuestas : dict
        Diccionario que mapea las respuestas posibles (claves) a valores numéricos (valores). 
        Ejemplo: {"Muy de acuerdo": 5, "De acuerdo": 4, ...}

    Proceso de la función
    ---------------------
    1. Se inicializa un diccionario vacío (conteo_respuestas) para acumular los conteos.
    2. Se extraen las listas de respuestas y valores del diccionario mapa_respuestas:
       - respuestas_posibles (lista de claves).
       - valores_posibles (lista de valores numéricos).
    3. Se recorren las columnas del DataFrame (enumerate(dataframe.columns, start=1)):
       - 'i' representará el índice o número de pregunta.
       - 'pregunta' será el nombre de la columna.
       - Para cada pregunta, se obtiene un conteo de respuestas usando `dataframe[pregunta].value_counts()`.
    4. Se rellena el diccionario de conteos con cada respuesta posible para esa pregunta:
       - Por cada par (respuesta, valor) en zip(respuestas_posibles, valores_posibles):
         - Se construye la clave: f"PREGUNTA_{i}_{valor}".
         - Se asigna el número de ocurrencias de dicha 'respuesta' en la serie `conteo`. 
           Si la respuesta no existe en la columna, se asigna 0.

    Valor de retorno
    ----------------
    dict
        Diccionario que asocia claves en el formato 'PREGUNTA_X_Y' 
        (X = número de pregunta, Y = valor numérico de la respuesta) 
        con el conteo de cuántas veces apareció esa respuesta en dicha pregunta.
        Ejemplo de clave: "PREGUNTA_3_5".

    Ejemplo de uso
    --------------
    >>> import pandas as pd
    >>> df = pd.DataFrame({
    ...     "Pregunta 1": ["Muy de acuerdo", "De acuerdo", "Muy de acuerdo"],
    ...     "Pregunta 2": ["En desacuerdo", "Muy de acuerdo", "De acuerdo"]
    ... })
    >>> mapa = {"Muy de acuerdo": 5, "De acuerdo": 4, "Ni de acuerdo ni en desacuerdo": 3,
    ...         "En desacuerdo": 2, "Muy en desacuerdo": 1}
    >>> conteos = obtenerRespuestas(df, mapa)
    >>> print(conteos)
    {
       "PREGUNTA_1_5": 2,
       "PREGUNTA_1_4": 1,
       "PREGUNTA_1_3": 0,
       "PREGUNTA_1_2": 0,
       "PREGUNTA_1_1": 0,
       "PREGUNTA_2_5": 1,
       "PREGUNTA_2_4": 1,
       "PREGUNTA_2_3": 0,
       "PREGUNTA_2_2": 1,
       "PREGUNTA_2_1": 0
    }
    """
    conteo_respuestas = {}
    valores_posibles = range(inicio, fin)
    
    for i, pregunta in enumerate(dataframe.columns, start=1):
        # Contar respuestas para la pregunta
        conteo = dataframe[pregunta].value_counts()
        
        for valor in valores_posibles:
            clave = f"PREGUNTA_{i}_{valor}"
            conteo_respuestas[clave] = conteo.get(valor, 0)  # Obtener el conteo o 0 si no aparece
    
    return conteo_respuestas

def calcularValores(respuestas_agrupadas):
    """
    Calcula la media y la desviación estándar para tres columnas clave de un DataFrame: 
    'Satisfaccion_Intrinseca', 'Satisfaccion_Extrinseca' y 'Satisfaccion_General'.

    Parámetros
    ----------
    respuestas_agrupadas : pd.DataFrame
        DataFrame que contiene las columnas anteriores con valores numéricos.

    Retorna
    -------
    dict
        Diccionario con las claves:
        - "MEDIA_INTRINSECA", "MEDIA_EXTRINSECA", "MEDIA_GENERAL" (medias),
        - "STD_INTRINSECA", "STD_EXTRINSECA", "STD_GENERAL" (desviaciones),
        todas redondeadas a 2 decimales.
    """
    reemplazos = {
        "MEDIA_INTRINSECA": round(respuestas_agrupadas['Satisfaccion_Intrinseca'].mean(), 2),
        "MEDIA_EXTRINSECA": round(respuestas_agrupadas['Satisfaccion_Extrinseca'].mean(), 2),
        "MEDIA_GENERAL": round(respuestas_agrupadas['Satisfaccion_General'].mean(), 2),
        "STD_INTRINSECA": round(respuestas_agrupadas['Satisfaccion_Intrinseca'].std(), 2),
        "STD_EXTRINSECA": round(respuestas_agrupadas['Satisfaccion_Extrinseca'].std(), 2),
        "STD_GENERAL": round(respuestas_agrupadas['Satisfaccion_General'].std(), 2),
    }

    return reemplazos

def escogerMedidas(media, archivo_medidas):
    """
    Carga los datos de rangos y medidas desde medidas.json,
    clasifica la media y devuelve un diccionario con 3 medidas seleccionadas
    aleatoriamente en función del nivel obtenido.
    """
    # Carga de datos desde el archivo JSON
    with open(archivo_medidas, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    #print(data)
    # Extracción de los rangos y las medidas
    rangos = data["rangos"]
    medidas = data["medidas"]

    # Clasificación según la media
    if rangos["rojo"][0] <= media <= rangos["rojo"][1]:
        nivel = "rojo"
    elif rangos["naranja"][0] <= media <= rangos["naranja"][1]:
        nivel = "naranja"
    elif rangos["amarillo"][0] <= media <= rangos["amarillo"][1]:
        nivel = "amarillo"
    elif rangos["verde"][0] <= media <= rangos["verde"][1]:
        nivel = "verde"
    else:
        nivel = None
    
    # Selección de medidas
    generar = 3
    if nivel:
        
        medidas_seleccionadas = random.sample(medidas[nivel], generar)
        return {
            "Prueba": nivel,
            "MEDIDAS": generar,
            "MEDIDA_1": medidas_seleccionadas[0],
            "MEDIDA_2": medidas_seleccionadas[1],
            "MEDIDA_3": medidas_seleccionadas[2]
        }
    else:
        return {
            "nivel": "Fuera de rango",
            "MEDIDAS": generar,
            "MEDIDA_1": "",
            "MEDIDA_2": "",
            "MEDIDA_3": ""
        }


def generarWord(plantilla_doc, carpeta_informes, reemplazos):
    """
    Crea un documento de Word a partir de una plantilla, reemplazando cada marcador 
    (clave) del diccionario `reemplazos` por su valor correspondiente.

    Parámetros
    ----------
    plantilla_doc : str
        Ruta al archivo .docx que sirve de plantilla.
    reemplazos : dict
        Diccionario cuyas claves son nombres de marcador y cuyos valores son 
        los textos que se insertarán en dichos marcadores.

    Genera
    ------
    Informe_Satisfaccion_Generado.docx
        Un archivo de Word con todos los marcadores reemplazados.
    """
    # Carpeta para guardar los informes
    if not os.path.exists(carpeta_informes):
        os.makedirs(carpeta_informes)
    
    doc = Document(plantilla_doc)

    # Aplicar reemplazos usando map
    list(map(lambda pair: replace_bookmark_pair(doc, pair), reemplazos.items()))

    output_doc = os.path.join(carpeta_informes, f"Informe_Satisfaccion_{reemplazos['NOMBRE_EMPRESA']}.docx")
    doc.save(output_doc)
    
    print(f"Informe generdo correctamente. Cierre esta ventana y vaya a {output_doc}")

    from io import BytesIO
from docx import Document

def generarWord_bytes(plantilla_doc: str, informe: str, reemplazos: dict) -> bytes:
    doc = Document(plantilla_doc)

    list(map(lambda pair: replace_bookmark_pair(doc, pair), reemplazos.items()))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generar_informe_satisfaccion(csv_source, empresa, invitados, num_medidas=3) -> bytes:
    ruta_script = os.path.dirname("./Satisfacción laboral/")
    carpeta_plantillas = os.path.join(ruta_script, "Plantillas")
    ruta_info_prl = os.path.join(ruta_script, "informacion_prl.json")
    plantilla_path = os.path.join(carpeta_plantillas, "plantilla_satisfaccion_laboral.docx")
    archivo_medidas = os.path.join(ruta_script, "medidas.json")

    if hasattr(csv_source, "read"):
        respuestas = pd.read_csv(csv_source, sep=None, engine="python")
    else:
        respuestas = pd.read_csv(csv_source, sep=None, engine="python")

    informacion = {
        "NOMBRE_EMPRESA": empresa,
        "PARTICIPACION": round(len(respuestas) / invitados * 100, 2) if invitados > 0 else 0
    }

    try:
        with open(ruta_info_prl, "r", encoding="utf-8") as f:
            informacion.update(json.load(f))
    except FileNotFoundError:
        pass

    # Obtener las preguntas directamente de las cabeceras del CSV
    preguntas = list(respuestas.columns)

    # Separar las preguntas en intrínsecas (pares) y extrínsecas (impares)
    preguntas_intrinsecas = [q for i, q in enumerate(preguntas) if (i + 1) % 2 == 0]
    preguntas_extrinsecas = [q for i, q in enumerate(preguntas) if (i + 1) % 2 != 0]

    # Mapear respuestas textuales a valores numéricos
    mapa_respuestas = {
        "Muy insatisfecho": 1,
        "Insatisfecho": 2,
        "Moderadamente insatisfecho": 3,
        "Ni satisfecho ni insatisfecho": 4,
        "Moderadamente satisfecho": 5,
        "Satisfecho": 6,
        "Muy satisfecho": 7
    }

    # Convertir respuestas textuales a numéricas usando el mapeo
    respuestas_convertidas = respuestas.copy()
    respuestas_convertidas = respuestas_convertidas.map(lambda x: mapa_respuestas.get(x, x) if x in mapa_respuestas else x)

    # Calcular las puntuaciones
    respuestas_agrupadas = pd.DataFrame()
    respuestas_agrupadas['Satisfaccion_Intrinseca'] = respuestas_convertidas[preguntas_intrinsecas].sum(axis=1)
    respuestas_agrupadas['Satisfaccion_Extrinseca'] = respuestas_convertidas[preguntas_extrinsecas].sum(axis=1)
    respuestas_agrupadas['Satisfaccion_General'] = respuestas_agrupadas['Satisfaccion_Intrinseca'] + respuestas_agrupadas['Satisfaccion_Extrinseca']

    calculos = calcularValores(respuestas_agrupadas)
    conteo_respuestas = obtenerRespuestas(respuestas, 1, 8)
    medidas = escogerMedidas(calculos['MEDIA_GENERAL'], archivo_medidas)

    reemplazos = informacion | calculos | conteo_respuestas | medidas

    doc = Document(plantilla_path)
    list(map(lambda pair: replace_bookmark_pair(doc, pair), reemplazos.items()))
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

