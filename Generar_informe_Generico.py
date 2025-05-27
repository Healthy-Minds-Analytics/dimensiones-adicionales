import os
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import glob
import json
import random
from copy import deepcopy
from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO

def seleccionar_csv(ruta):
    """Busca archivos CSV en la carpeta de la ruta proporcionada.
       Si solo hay uno, lo devuelve, y si hay más de uno
       escoge el más reciente."""
    
    patron_busqueda = os.path.join(ruta, "*.csv")
    archivos_csv = glob.glob(patron_busqueda)
    
    if not archivos_csv:
        print(f"No se encontró ningún archivo CSV en la carpeta {ruta}.")
        return None

    if len(archivos_csv) == 1:
        return archivos_csv[0]

    # Si hay varios, escogemos el más reciente (mayor fecha de modificación)
    archivo_mas_reciente = max(archivos_csv, key=os.path.getmtime)
    return archivo_mas_reciente

def replace_bookmark_pair(doc, pair):
    """
    Reemplaza el contenido asociado a un marcador específico en un documento Word (python-docx),
    recorriendo todo su árbol XML, incluidos cuadros de texto y demás estructuras anidadas.

    Parámetros
    ----------
    doc : docx.Document
        Objeto Document proporcionado por la librería python-docx. Representa el documento 
        donde se realizará la búsqueda y reemplazo.
    pair : tuple
        Tupla que contiene (bookmark_name, replacement).
        
        - bookmark_name (str): Nombre del marcador a localizar en el documento.
        - replacement (str): Texto o valor que se asignará en sustitución del contenido 
          hallado dentro del marcador.

    Comportamiento
    -------------
    1. Se define una función interna `replace_in_element(element)` que:
       - Recorre recursivamente cada uno de los subelementos del XML del documento.
       - Si encuentra un 'bookmarkStart' cuyo atributo 'w:name' coincida con bookmark_name:
         - Marca la variable `found` como True.
         - Avanza sobre los elementos hermanos (next_sibling) del marcador hasta localizar un 
           run (`<w:r>`) que contenga un elemento texto (`<w:t>`).
         - Reemplaza el contenido de `<w:t>` con la cadena `replacement`.
         - Luego elimina (en caso de existir) todos los elementos hermanos siguientes 
           hasta toparse con un 'bookmarkEnd' (indica el fin del marcador).
         - Termina el proceso tras el primer reemplazo exitoso.
       - Continúa explorando recursivamente el resto de elementos si no se ha encontrado el marcador.

    2. La función principal `replace_bookmark_pair(doc, pair)`:
       - Toma la raíz (`doc._element`) y la recorre llamando a `replace_in_element`.
       - Si, al finalizar el recorrido, la variable `found` sigue en False, 
         imprime un aviso por consola indicando que el marcador no se encontró.

    Notas
    ----
    - Este método modifica el documento en memoria: al finalizar, conviene llamar a `doc.save(...)`
      para persistir los cambios en un archivo.
    - La función solo realiza un reemplazo por marcador. Si un marcador aparece varias veces 
      con el mismo nombre, solo se reemplazará la primera aparición que se halle al recorrer el XML.
    - El proceso recursivo permite hallar el marcador aunque esté dentro de cuadros de texto, 
      tablas u otras secciones anidadas del documento.

    Ejemplo de uso
    --------------
    >>> from docx import Document
    >>> doc = Document("mi_documento.docx")
    >>> replace_bookmark_pair(doc, ("MI_MARKER", "Nuevo contenido"))
    >>> doc.save("mi_documento_modificado.docx")
    """

    bookmark_name, replacement = pair
    found = False

    def replace_in_element(element):
        nonlocal found
        for child in element:
            if child.tag.endswith('bookmarkStart') and child.get(qn('w:name')) == bookmark_name:
                found = True

                # Encuentra el <w:r> que contiene <w:t>
                run_elem = child.getnext()
                while run_elem is not None and not run_elem.tag.endswith('r'):
                    run_elem = run_elem.getnext()
                if run_elem is None:
                    return

                # Localiza la etiqueta <w:t>
                text_elem = run_elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                if text_elem is None:
                    return

                # Obtiene su párrafo padre (<w:p>)
                p_elem = run_elem.getparent()

                # Divide el replacement en líneas
                lines = str(replacement).split('\n')
                text_elem.text = lines[0]

                # Para cada línea extra, crea un párrafo nuevo después de p_elem
                for line in lines[1:]:
                    # clonamos el nodo <w:p> completo (incluye <w:pPr> con bullet o numeración)
                    new_p = deepcopy(p_elem)
                    # dentro de ese párrafo clonado, vaciamos todos los runs
                    for r in new_p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                        new_p.remove(r)
                    # creamos un run nuevo con el texto
                    new_r = OxmlElement('w:r')
                    new_t = OxmlElement('w:t')
                    new_t.text = line
                    new_r.append(new_t)
                    new_p.append(new_r)
                    # insertamos este nuevo párrafo justo después del original
                    p_elem.addnext(new_p)
                    # avanzamos el cursor para la siguiente iteración
                    p_elem = new_p

                return

            replace_in_element(child)

    replace_in_element(doc._element)
    if not found:
        print(f"Marcador '{bookmark_name}' no encontrado")

def obtenerRespuestas(dataframe, inicio, fin):
    """
    Genera un diccionario con el conteo de cada respuesta por pregunta en un DataFrame, 
    construyendo las claves en el formato 'PREGUNTA_X_Y'.

    Parámetros
    ----------
    dataframe : pd.DataFrame
        DataFrame donde cada columna representa una pregunta y cada fila registra las 
        respuestas de un encuestado. Los valores del DataFrame deben coincidir con las 
        posibles respuestas definidas en mapa_respuestas.
    mapa_respuestas : dict
        Diccionario que mapea las respuestas posibles (claves) a valores numéricos (valores). 
        Ejemplo: {"Muy de acuerdo": 5, "De acuerdo": 4, ...}

    Proceso de la función
    ---------------------
    1. Se inicializa un diccionario vacío (conteo_respuestas) para acumular los conteos.
    2. Se extraen las listas de respuestas y valores del diccionario mapa_respuestas:
       - respuestas_posibles (lista de claves).
       - valores_posibles (lista de valores numéricos).
    3. Se recorren las columnas del DataFrame (enumerate(dataframe.columns, start=1)):
       - 'i' representará el índice o número de pregunta.
       - 'pregunta' será el nombre de la columna.
       - Para cada pregunta, se obtiene un conteo de respuestas usando `dataframe[pregunta].value_counts()`.
    4. Se rellena el diccionario de conteos con cada respuesta posible para esa pregunta:
       - Por cada par (respuesta, valor) en zip(respuestas_posibles, valores_posibles):
         - Se construye la clave: f"PREGUNTA_{i}_{valor}".
         - Se asigna el número de ocurrencias de dicha 'respuesta' en la serie `conteo`. 
           Si la respuesta no existe en la columna, se asigna 0.

    Valor de retorno
    ----------------
    dict
        Diccionario que asocia claves en el formato 'PREGUNTA_X_Y' 
        (X = número de pregunta, Y = valor numérico de la respuesta) 
        con el conteo de cuántas veces apareció esa respuesta en dicha pregunta.
        Ejemplo de clave: "PREGUNTA_3_5".

    Ejemplo de uso
    --------------
    >>> import pandas as pd
    >>> df = pd.DataFrame({
    ...     "Pregunta 1": ["Muy de acuerdo", "De acuerdo", "Muy de acuerdo"],
    ...     "Pregunta 2": ["En desacuerdo", "Muy de acuerdo", "De acuerdo"]
    ... })
    >>> mapa = {"Muy de acuerdo": 5, "De acuerdo": 4, "Ni de acuerdo ni en desacuerdo": 3,
    ...         "En desacuerdo": 2, "Muy en desacuerdo": 1}
    >>> conteos = obtenerRespuestas(df, mapa)
    >>> print(conteos)
    {
       "PREGUNTA_1_5": 2,
       "PREGUNTA_1_4": 1,
       "PREGUNTA_1_3": 0,
       "PREGUNTA_1_2": 0,
       "PREGUNTA_1_1": 0,
       "PREGUNTA_2_5": 1,
       "PREGUNTA_2_4": 1,
       "PREGUNTA_2_3": 0,
       "PREGUNTA_2_2": 1,
       "PREGUNTA_2_1": 0
    }
    """
    conteo_respuestas = {}
    valores_posibles = range(inicio, fin)
    
    for i, pregunta in enumerate(dataframe.columns, start=1):
        # Contar respuestas para la pregunta
        conteo = dataframe[pregunta].value_counts()
        
        for valor in valores_posibles:
            clave = f"PREGUNTA_{i}_{valor}"
            conteo_respuestas[clave] = conteo.get(valor, 0)  # Obtener el conteo o 0 si no aparece
    
    return conteo_respuestas

def calcularValores(respuestas_dim: pd.DataFrame) -> pd.DataFrame:
    """
    Devuelve un DataFrame con la media y la desviación estándar
    por cada dimensión (como filas), ya redondeadas a 2 decimales.
    
    Índices: nombre de la dimensión.
    Columnas: ['mean', 'std'].
    """
    stats = respuestas_dim.agg(['mean', 'std']).T
    return stats.round(2)

def df_a_reemplazos(df_stats: pd.DataFrame) -> dict:
    reemplazos = {}
    for dim, row in df_stats.iterrows():
        reemplazos[f"MEDIA_{dim}"] = row['mean']
        reemplazos[f"STD_{dim}"] = row['std']
    return reemplazos

def escogerMedidas(estadisticas: pd.DataFrame, limite=10) -> dict:
    estadisticas_corregidas = estadisticas.copy()
    estadisticas_corregidas.loc[['FISICAS', 'SOCIALES', 'PSICOLOGICAS'], 'mean'] *= 3

    # Buscar las dimensiones (filas) de estadisticas cuya media sea mayor que el límite indicado
    alertas = estadisticas_corregidas[estadisticas_corregidas['mean'] > limite]

    dims = list(alertas.index)

    if len(dims) < 2:
        dims = estadisticas_corregidas['mean'].nlargest(2).index.tolist()


    ficheros = {
        "CARACTERISTICAS_TAREA": 'caracteristicas_tarea',
        "ORGANIZACION":'organizacion',
        "TEDIO": 'tedio',
        "CANSANCIO_EMOCIONAL": 'cansancio_emocional',
        "DESPERSONALIZACION": 'despersonalizacion',
        "REALIZACION_PERSONAL": 'realizacion_personal',
        'FISICAS': 'consecuencias_fisicas',
        'SOCIALES': 'consecuencias_sociales',
        'PSICOLOGICAS': 'consecuencias_psicologicas'
        
    }

    parrafos = []
    for dim in dims:
        base = ficheros.get(dim, '')
        if not base:
            # no hay JSON asociado: lo ignoramos
            print(f"No hay fichero asociado para la dimensión {dim!r}, por lo que no se proponen medidas para dicha dimensión.")
            continue

        fichero = os.path.join('Medidas', f'{base}.json')
        if not os.path.exists(fichero):
            print(f"El fichero {fichero} no existe")
            continue
        
        # Abrir archivo JSON
        with open(fichero, 'r', encoding='utf-8') as f:
            data = json.load(f)
        if not data:
            print(f"El fichero {fichero} de la dimensión {dim!r} está vacío")
            continue
        
        try:
            lista_medidas = next(iter(data.values()))
        except StopIteration:
            continue

        if not lista_medidas:
            continue

        #print(lista_medidas)
        # Escoge una medida al azar
        medida = random.choice(lista_medidas)

        parrafos.append(f"{medida}")

    texto_final = "\n".join(parrafos)

    return {'MEDIDAS': texto_final}


def generarWord(plantilla_doc, informe, carpeta_informes, reemplazos):
    """
    Crea un documento de Word a partir de una plantilla, reemplazando cada marcador 
    (clave) del diccionario `reemplazos` por su valor correspondiente.

    Parámetros
    ----------
    plantilla_doc : str
        Ruta al archivo .docx que sirve de plantilla.
    informe : dict
        Nombre del informe a generar.
    carpeta_informes : str
        Ruta a la carpeta donde se guardará el informe.
    reemplazos : dict
        Diccionario cuyas claves son nombres de marcador y cuyos valores son 
        los textos que se insertarán en dichos marcadores.

    Genera
    ------
    Informe_Satisfaccion_Generado.docx
        Un archivo de Word con todos los marcadores reemplazados.
    """
    # Carpeta para guardar los informes
    if not os.path.exists(carpeta_informes):
        os.makedirs(carpeta_informes)
    
    doc = Document(plantilla_doc)

    # Aplicar reemplazos usando map
    list(map(lambda pair: replace_bookmark_pair(doc, pair), reemplazos.items()))

    output_doc = os.path.join(carpeta_informes, f"Informe_Burnout_{reemplazos['NOMBRE_EMPRESA']}.docx")
    doc.save(output_doc)
    
    print(f"Informe generdo correctamente. Cierre esta ventana y vaya a {output_doc}")

def calcular_metricas_estadisticas(respuestas_agrupadas: pd.DataFrame) -> dict:
    """
    Calcula estadísticas para cada dimensión y devuelve un diccionario de diccionarios.
    
    Estructura del retorno:
    {
        "Satisfaccion_General": {
            "media": ...,
            "std": ...,
            ...
        },
        ...
    }
    """
    metricas = {}

    dimensiones = [
        "Satisfaccion_General",
        "Satisfaccion_Intrinseca",
        "Satisfaccion_Extrinseca"
    ]

    for dimension in dimensiones:
        serie = respuestas_agrupadas[dimension]

        metricas[dimension] = {
            "media": round(serie.mean(), 2),
            "std": round(serie.std(), 2),
            "min": int(serie.min()),
            "max": int(serie.max()),
            "p25": round(serie.quantile(0.25), 2),
            "p75": round(serie.quantile(0.75), 2),
            "count": int(serie.count())
        }

    return metricas

def load_questions(json_data: dict, locale: str = "es") -> list[dict]:
    """
    Filtra y formatea preguntas a partir de un JSON ya cargado en memoria.

    Parámetros
    ----------
    json_data : dict
        Objeto JSON cargado que contiene 'availableLocales' y 'questions'.
    locale : str
        Código de idioma a seleccionar.

    Retorna
    -------
    List[Dict]
        Lista de preguntas con {'text', 'options'} para el locale.
    """
    # Validación de locale
    locales = json_data.get("availableLocales", [])
    if locale not in locales:
        raise ValueError(
            f"Locale '{locale}' no disponible. Solo están: {locales}"
        )
    preguntas = []
    for q in json_data.get("questions", []):
        texto = q.get("questionTexts", {}).get(locale, "").strip()
        opts = []
        for opt in q.get("options", []):
            opt_text = opt.get("optionTexts", {}).get(locale, "").strip()
            opts.append({
                "text": opt_text,
                "value": opt.get("value")
            })
        preguntas.append({"text": texto, "options": opts})
    return preguntas

def insert_paragraph_after(paragraph, text=None, style=None):
    """
    Inserta un nuevo párrafo justo después de `paragraph` dado.
    Devuelve el objeto Paragraph creado.
    """
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para

def generar_informe_generico(csv_source, json_source, empresa: str, titulo: str, invitados: int, locale: str = "es") -> bytes:
    """
    Genera un informe genérico leyendo:
      - csv_source: ruta o UploadedFile de Streamlit con las respuestas.
      - json_source: ruta o UploadedFile de Streamlit con las preguntas.
    Devuelve el .docx en memoria (bytes) listo para descargar.
    """
    # 1) Plantilla
    ruta_script = os.path.dirname("./Generico/")
    carpeta_plantillas = os.path.join(ruta_script, "Plantillas")
    plantilla_path = os.path.join(carpeta_plantillas, "plantilla_generico.docx")

    # 2) Carga del JSON de preguntas (file-like o ruta)
    if isinstance(json_source, dict):
        json_data = json_source
    elif hasattr(json_source, "read"):
        json_source.seek(0)
        json_data = json.load(json_source)
    else:
        json_data = json.load(open(json_source, encoding="utf-8"))

    preguntas = load_questions(json_data, locale)

    # 3) Lectura del CSV
    if hasattr(csv_source, "read"):
        respuestas = pd.read_csv(csv_source, sep=";", engine="python")
    else:
        respuestas = pd.read_csv(csv_source, sep=";", engine="python")

    
    # 4) Info fija
    info = {
        "NOMBRE_EMPRESA": empresa,
        "TITULO_INFORME":  titulo,
        "PARTICIPACION":   round(len(respuestas) / invitados * 100, 2) if invitados>0 else 0
    }

    # 5) Mapeo texto→valor
    mapa_respuestas = {'no':0, 'sí':10, 'si':10}
    df_val = respuestas.applymap(
        lambda x: mapa_respuestas.get(str(x).strip().lower(),
                     int(x) if str(x).isdigit() else x)
        if isinstance(x, str) else x
    )

    # 6) Conteos y stats
    from Generar_informe_Generico import obtenerRespuestas, calcularValores  # o importa desde utils
    conteos  = obtenerRespuestas(df_val, inicio=1, fin=11)
    df_stats = calcularValores(df_val)

    # 7) Montaje del DOCX en memoria
    doc = Document(plantilla_path)

    # 7.1) Reemplazo de marcadores con info fija
    for marcador, texto in info.items():
        replace_bookmark_pair(doc, (marcador, str(texto)))

    # 7.2) Inserción dinámica de preguntas y resultados
    # Busca párrafo-ancla
    for p in doc.paragraphs:
        if "TEXTO_PREGUNTAS" in p.text:
            anchor = p
            anchor.text = anchor.text.replace("TEXTO_PREGUNTAS", "")
            break
    else:
        raise RuntimeError("Marcador TEXTO_PREGUNTAS no encontrado")

    current = anchor
    for idx, q in enumerate(preguntas, start=1):
        # pregunta
        print(f"Imprimiendo {q}")
        current = insert_paragraph_after(current, q["text"], style="Normal")
        # opciones y conteos
        for opt in q["options"]:
            raw = opt["value"]
            val = mapa_respuestas.get(str(raw).strip().lower(), 
                                      int(raw) if str(raw).isdigit() else None)
            cnt = conteos.get(f"PREGUNTA_{idx}_{val}", 0)
            current = insert_paragraph_after(current, f"{opt['text']}: {cnt}", style="Bullet list")
        # estadísticos
        stats = df_stats.iloc[idx-1]
        current = insert_paragraph_after(current, "Resultados:", style="Normal")
        for etiqueta, campo in (("Media","mean"),("Desviación típica","std")):
            current = insert_paragraph_after(
                current,
                f"{etiqueta}: {stats[campo]:.2f}",
                style="Bullet list"
            )

    # 8) Volcado a bytes
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()